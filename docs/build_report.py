"""Build PROJECT_REPORT.docx — embeds the five screenshots into a styled Word
document built from the same source-of-truth as PROJECT_REPORT.md.

Run from the project root:  python docs/build_report.py
Output:                     docs/PROJECT_REPORT.docx
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# --- paths ------------------------------------------------------------------
HERE = Path(__file__).parent
SHOTS = HERE / 'screenshots'
OUT = HERE / 'PROJECT_REPORT.docx'


def _resolve_out_path(preferred: Path) -> Path:
    """If Word is holding the file open, fall back to a numbered sibling so
    the build never silently fails for the user."""
    try:
        if preferred.exists():
            with open(preferred, 'a'):
                pass
        return preferred
    except PermissionError:
        for i in range(2, 99):
            alt = preferred.with_name(f'{preferred.stem}_v{i}{preferred.suffix}')
            try:
                if alt.exists():
                    with open(alt, 'a'):
                        pass
                return alt
            except PermissionError:
                continue
        raise

# --- colours ----------------------------------------------------------------
NAVY = RGBColor(0x10, 0x2A, 0x43)
TEAL = RGBColor(0x06, 0x76, 0x8A)
SLATE = RGBColor(0x33, 0x41, 0x55)
BODY = RGBColor(0x1F, 0x29, 0x37)
MUTED = RGBColor(0x6B, 0x72, 0x80)
CRIT = RGBColor(0xC0, 0x39, 0x2B)
WARN = RGBColor(0xC9, 0x8A, 0x0C)
INFO = RGBColor(0x42, 0x6B, 0xA5)


# --- helpers ----------------------------------------------------------------
def _cell_shading(cell, hex_fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_fill)
    tc_pr.append(shd)


def set_default_font(doc, name='Calibri', size=11):
    style = doc.styles['Normal']
    style.font.name = name
    style.font.size = Pt(size)
    style.font.color.rgb = BODY


def title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.font.size = Pt(28)
    r.font.bold = True
    r.font.color.rgb = NAVY


def subtitle(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(13)
    r.italic = True
    r.font.color.rgb = TEAL


def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(18)
    r.font.bold = True
    r.font.color.rgb = NAVY


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = TEAL


def h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = SLATE


def para(doc, text, italic=False, bold=False, muted=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    if muted:
        r.font.color.rgb = MUTED
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(text)


def numbered(doc, text):
    p = doc.add_paragraph(style='List Number')
    p.add_run(text)


def divider(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'C5CDD8')
    pBdr.append(bottom)
    pPr.append(pBdr)


def image(doc, fname, caption=None, width_in=6.4):
    path = SHOTS / fname
    if not path.exists():
        para(doc, f'[missing screenshot: {fname}]', italic=True, muted=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width_in))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(f'Figure — {caption}')
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = MUTED


def table(doc, headers, rows, col_widths=None, severity_col=None):
    """Render a styled table. severity_col=int colours the cell text per row."""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.autofit = False
    t.style = 'Light Grid Accent 1'
    # header
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _cell_shading(hdr[i], '102A43')
    # body
    for ri, row in enumerate(rows, start=1):
        cells = t.rows[ri].cells
        for ci, val in enumerate(row):
            cells[ci].text = ''
            p = cells[ci].paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(10)
            if severity_col is not None and ci == severity_col:
                sev = str(val).upper()
                if sev == 'CRITICAL':
                    r.font.color.rgb = CRIT
                    r.bold = True
                elif sev == 'WARNING':
                    r.font.color.rgb = WARN
                    r.bold = True
                elif sev == 'INFO':
                    r.font.color.rgb = INFO
    # column widths
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    return t


def code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(text)
    r.font.name = 'Consolas'
    r.font.size = Pt(9)
    r.font.color.rgb = SLATE
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F5F7FA')
    pPr.append(shd)


def page_break(doc):
    doc.add_page_break()


def lead(text):
    """Bold teal label that introduces the paragraph beneath it."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = TEAL


def rule_card(rule_id, title, severity, what, pass_ex, fail_ex, why):
    """Render one rule's detailed walkthrough."""
    # Title line: 'GEN-01 — Streaming silence > 5 min'  with severity chip
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f'{rule_id}  ')
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = SLATE
    r2 = p.add_run(title)
    r2.bold = True
    r2.font.size = Pt(11)
    r2.font.color.rgb = NAVY
    r3 = p.add_run(f'  [{severity}]')
    r3.bold = True
    r3.font.size = Pt(9)
    if severity == 'CRITICAL':
        r3.font.color.rgb = CRIT
    elif severity == 'WARNING':
        r3.font.color.rgb = WARN
    else:
        r3.font.color.rgb = INFO

    def _line(label, text, mono=False):
        pp = doc.add_paragraph()
        pp.paragraph_format.space_after = Pt(1)
        pp.paragraph_format.left_indent = Cm(0.3)
        ra = pp.add_run(f'{label} ')
        ra.bold = True
        ra.font.size = Pt(10)
        ra.font.color.rgb = TEAL
        rb = pp.add_run(text)
        rb.font.size = Pt(10)
        if mono:
            rb.font.name = 'Consolas'
            rb.font.color.rgb = SLATE

    _line('What it checks.', what)
    _line('Passes when.', pass_ex)
    _line('Fails when.', fail_ex)
    _line('Why it matters.', why)


# ---------------------------------------------------------------------------
# Rule walkthrough data
# ---------------------------------------------------------------------------
GEN_RULES = [
    ('GEN-01', 'Streaming silence > 5 min during production hours', 'WARNING',
     'A tag must keep publishing during production hours (06:00–22:00 IST). Any silence longer than 5 min mid-shift is flagged as a comms drop.',
     'Reading at 14:32:15, next at 14:33:14 — gap 59 s. OK.',
     'Reading at 14:32:15, next at 14:42:18 — gap 10 m 3 s during shift. Emits "Streaming stopped 10m3s".',
     'Silent sensors are the most common pre-failure signal. Catching them within minutes is the difference between a hot fix and a halted line.'),
    ('GEN-02', 'Null or empty value', 'WARNING',
     'A row\'s Value must not be null or whitespace. Bag-out idle (a single ",") is exempt because that is a legitimate "no bag in dispenser" state.',
     'Value = "D:202641,S:3,B:38,..." — OK.',
     'Value = "" — emits "Null or empty value".',
     'A null on the wire usually means the PLC tag was unbound or the read failed; the value cannot be trusted by any downstream rule.'),
    ('GEN-03', 'Timestamps must be monotonically increasing per tag', 'INFO',
     'Successive readings on the same tag must have non-decreasing timestamps.',
     'T1 = 14:32:15, T2 = 14:33:14 — OK.',
     'T1 = 14:33:14, T2 = 14:32:15 — emits "Out-of-order TS".',
     'Out-of-order timestamps indicate gateway clock skew, message replay, or buffer flushes that distort every downstream time-series calculation.'),
    ('GEN-04', 'Source metadata sanity', 'INFO',
     'Every row must have SiteId="LLPL", IotDeviceId="uaq-lakme-hul-iotedge-01", SensorId="opcua".',
     'All three match — OK.',
     'SiteId = "HUL" — emits "Unexpected source metadata: SiteId=HUL".',
     'Catches mis-routed telemetry from another plant or a swapped gateway before it contaminates the dashboard.'),
]

PSM_RULES = [
    ('PSM-01', 'Schema must contain all of D, S, B, R, RM, SP, PV', 'CRITICAL',
     'Every PSM batch string must carry the full seven-field schema.',
     '"D:202641,S:3,B:38,R:PLUMERIA NOODLES,RM:EDTA,SP:2.0499,PV:2.0850" — OK.',
     '"D:202641,S:3,B:38,R:PLUMERIA,RM:EDTA,SP:2.0499" — emits "Schema incomplete: PV".',
     'All downstream PSM rules depend on these fields; a missing field invalidates every analysis run against the row.'),
    ('PSM-02', 'PV must not be 0 while status = dosing (S=2)', 'CRITICAL',
     'When status code S = 2 (actively dosing), the process value must register some weight (≥ 0.001 kg).',
     'S=2, PV=1.45 kg — dosing is visible, OK.',
     'S=2, PV=0.0 — emits "PV0 while dosing. PV=0".',
     'A scale reading zero while the dosing valve is open means the flow sensor is stuck or the material is blocked — both stop production.'),
    ('PSM-03', 'PV must not be negative', 'CRITICAL',
     'PV must be ≥ 0. Weight is a physical quantity; negative values are impossible.',
     'PV = 2.085 kg — OK.',
     'PV = −0.123 — emits "PV is negative (-0.123)".',
     'Negative weight indicates a tare error or calibration drift; the affected scale is unreliable until recalibrated.'),
    ('PSM-04', 'PV must not drop > 0.5 kg within the same batch counter', 'WARNING',
     'Within the same batch counter B, PV must not drop by more than 0.5 kg between consecutive polls. Dosing is monotonic by design.',
     'B=38, PV history 1.2 → 1.5 → 1.8 kg — monotonic, OK.',
     'B=38, PV history 2.0 → 1.4 kg — emits "PV dropped within batch 38 (2->1.4)".',
     'A drop mid-batch means the valve flapped, the tank leaked, or the sensor glitched — all need immediate investigation.'),
    ('PSM-05', 'Completion deviation |PV−SP|/SP ≤ 5 % when batch is completed', 'WARNING',
     'At completion (S=3), the dosing error |PV − SP| / SP must be ≤ 5 %.',
     'S=3, SP=2.05 kg, PV=2.08 kg — 1.7 % deviation, OK.',
     'S=3, SP=35.74 kg, PV=38.25 kg — 7.0 % — emits "Completion deviation 7.0% > 5%".',
     '5 % is the per-RM yield tolerance. Anything beyond means the batch is off-spec and may need to be scrapped or reworked.'),
    ('PSM-06', 'Status code must be 1, 2, or 3', 'WARNING',
     'Status code S must be exactly 1 (idle), 2 (dosing), or 3 (complete).',
     'S=2 — OK.',
     'S=5 — emits "Invalid status code S=5".',
     'Garbage status codes signal corrupted polls or firmware regressions; the row\'s status cannot be trusted by any subsequent rule.'),
    ('PSM-07', 'Batch counter sequential or 39→0 wrap', 'WARNING',
     'Batch counter B should equal previous (same batch), previous + 1 (advance), or wrap from 39 back to 0.',
     'previousB=38, B=39 — advance, OK. previousB=39, B=0 — wrap, OK.',
     'previousB=38, B=42 — emits "Non-sequential batch jump 38->42".',
     'A jump means some polls were lost; the data for the missed batches is permanently unrecoverable.'),
    ('PSM-08', 'All RMs in current cycle agree on D and R', 'WARNING',
     'All raw materials being dosed simultaneously must agree on D (date) and R (recipe).',
     'AOS row D=202641,R=PLUMERIA; Caustic row D=202641,R=PLUMERIA — OK.',
     'AOS R=PLUMERIA; Caustic R=JASMINE — emits "Recipe/date mismatch across RMs".',
     'Two RMs reporting different recipes in the same cycle is an orchestration bug; the resulting batch is unusable.'),
    ('PSM-09', 'Streaming gap inside production hours (PSM-scoped)', 'INFO',
     'PSM-tag-scoped streaming gap. Same logic as GEN-01 but tagged PSM-09 so it surfaces under the PSM filter.',
     '60 s gap on a PSM tag — OK.',
     '10 min gap at 14:00 — emits "Comms gap 10m3s".',
     'PSM is the most rule-heavy zone; isolating its comms drops helps the PSM engineer triage faster than scanning the global GEN-01 stream.'),
    ('PSM-10', 'SP-change detection (event log only)', 'INFO',
     'Logs every change of SP between polls. Always passes — this is an event log, not a failure.',
     'SP=10.50 → SP=10.50 — silent OK.',
     'SP=10.50 → SP=10.55 — passes but emits "SP adjusted 10.5->10.55".',
     'Provides the audit trail for SP changes. This event log is the data source for Insight #9 (Recipe-drift detection in §11).'),
    ('PSM-11', 'Julian date D within 2 days of system time', 'INFO',
     'The Julian date D field (YYYYDDD format) must be within 2 days of the system clock.',
     'System day 41 of 2026; D = "202641" — OK.',
     'D = "202635" (6 days stale) — emits "D field stale by 6.0d".',
     'Stale D = PLC clock drift or message replay; alerts the team to a gateway-side problem before stale timestamps corrupt batch trace.'),
    ('PSM-12', 'Batch_PV_Weight ≈ Σ RM PVs (≤ 2 % deviation)', 'INFO',
     'The aggregate Batch_PV_Weight should match the sum of the individual RM PVs within 2 % — mass conservation across the dosing skid.',
     'Batch_PV_Weight=100 kg, Σ RM PV=99.5 kg — 0.5 % off, OK.',
     'Batch_PV_Weight=100 kg, Σ RM PV=92 kg — 8 % — emits "Batch weight vs RM-sum deviation 8.0%".',
     'A mass imbalance points to an unaccounted leak, a tare bug, or a missing RM row that the orchestrator forgot to record.'),
]

SMX_RULES = [
    ('SMX-13', 'Rework stuck > 0 for > 3 consecutive polls', 'WARNING',
     'The REWORK tag should sit at 0 in normal flow. A value > 0 for more than 3 consecutive polls indicates the mixer is stuck reworking material.',
     '0, 0, 30, 0 — one isolated rework event, OK.',
     '30, 30, 30, 30 (four consecutive non-zero polls) — emits "Rework stuck active for 4 polls (value=30)".',
     'Persistent rework points to a chronic mixing problem (e.g. lump formation, undersized impeller load). Material and downtime costs escalate quickly if not caught early.'),
    ('SMX-14', 'Rework value must be 0 or 30 (recognised codes)', 'WARNING',
     'The rework tag uses a discrete code set. Across the entire historian we only see two values: 0 (normal) and 30 (rework event). Anything else is treated as a corrupted reading.',
     'reworkValue=0 or 30 — OK.',
     'reworkValue=15 — emits "Unrecognised rework code 15 (expected 0 or 30)".',
     'Catches PLC firmware regressions or sensor glitches that emit values outside the discrete code set the line was designed around.'),
    ('SMX-15', '> 5 rework events in last 60 min — chronic rework', 'WARNING',
     'Tracks every non-zero rework event in a rolling 60-minute window per mixer. A burst of more than 5 events means the mixer is chronically reworking — beyond what SMX-13 catches on a single sticky run.',
     '2 events in last hour — OK.',
     '6 events in last hour — emits "6 rework events in last 60 min — chronic rework".',
     'Distinguishes a single stuck-rework episode (SMX-13) from a chronic-rework day. Drives the decision to halt the mixer for inspection vs. ride out one batch.'),
    ('SMX-16', 'MX<a> and MX<b> must not dose the same batch counter simultaneously', 'WARNING',
     'When two mixers are in status 2 (dosing) with the same batch counter B at the same time, they are feeding into one logical batch ID — a cross-contamination risk.',
     'MX1: S=2,B=38; MX2: S=2,B=39 — different batches, OK.',
     'MX1: S=2,B=38; MX2: S=2,B=38 — emits "MX1 & MX2 both dosing same batch 38".',
     'Material destined for two different batches gets blended together — guaranteed off-spec product on both.'),
    ('SMX-17', 'Recipe must not change mid-batch', 'INFO',
     'While the batch counter B is unchanged, the recipe R must not change.',
     'B=38, R=LMSU3R2_LBGRMEXP_600KG throughout — OK.',
     'B=38, R=LMSU3R2_...; then B=38, R=LTQA7R4_... — emits "Recipe changed mid-batch 38".',
     'Recipe should be locked when the batch starts. A mid-batch change means the orchestrator overwrote R without resetting B — guaranteed bad batch.'),
    ('SMX-18', 'Barcode must be empty/idle marker or pure numeric (legacy)', 'CRITICAL',
     'Legacy rule retained for backwards-compatibility with the old SM_MX*_BC scanner tags. New cascades expose a dedicated BATCHCOUNTER + RECIPE_NAME pair instead, so this rule fires only when a SIGMA_BARCODE row appears.',
     '"Scan Barcode" — OK (idle). "1300326026" — OK (valid scan).',
     '"ABC123XYZ" — emits "Malformed barcode scan \'ABC123XYZ\'".',
     'Keeps Sigma backwards-compatible with older cascades that still emit the legacy barcode tag.'),
]

SLO_RULES = [
    ('SLO-01', 'Noodle type must be in the valid set', 'CRITICAL',
     'The noodle type string must be one of: JASMINE, PLUMERIA, SERGIO 56, TEXAS MOD, GALAXY, 20 PKO TULIP, LILAC NOODLES.',
     '"JASMINE NOODLES" — OK.',
     '"PLUMARIA NOODLES" (typo) — emits "Unknown noodle type \'PLUMARIA NOODLES\'".',
     'Catches PLC label corruption or new noodle types that have not been onboarded into the validation set yet — easy to miss otherwise.'),
    ('SLO-02', 'Bag-out detail must be CSV of 4 fields, or single \',\' for idle', 'CRITICAL',
     'Bag-out value must be a 4-field CSV (batchId, SP, PV, noodleType), or a single "," meaning "no bag in dispenser".',
     '"B12345,850,847.2,JASMINE NOODLES" — OK. "," — idle, OK.',
     '"B12345,850" — only 2 fields — emits "Expected 4 CSV fields, got 2".',
     'Malformed bag-out CSV breaks every downstream consumer, including the bagging-station HMI that operators read.'),
    ('SLO-03', 'Bag PV weight ≥ 0', 'CRITICAL',
     'Bagging scale PV weight must be non-negative.',
     'PV = 850 kg — OK.',
     'PV = −50 — emits "Bag PV weight negative (-50)".',
     'Same physics as PSM-03 but for the bagging scale: a negative reading means the scale needs immediate recalibration.'),
    ('SLO-04', 'Bag PV within 10 % of SP', 'WARNING',
     'Bagging deviation tolerance is 10 % — wider than PSM\'s 5 % because bag scales are less precise than dosing scales.',
     'SP=850, PV=857 — 0.8 % off, OK.',
     'SP=850, PV=950 — 11.8 % — emits "Bag PV 950kg outside 10% of SP 850kg".',
     'Catches clear over- or under-fills before the bag reaches packaging and becomes scrap.'),
    ('SLO-05', 'Day-silo and Buffer-silo at same index agree on noodle type', 'WARNING',
     'Day_silo_N and Buffer_silo_N (same index) are physically paired and must report the same noodle type.',
     'Day_silo_3 = "PLUMERIA"; Buffer_silo_3 = "PLUMERIA" — OK.',
     'Day_silo_3 = "PLUMERIA"; Buffer_silo_3 = "JASMINE" — emits "Day/Buffer silo 3 mismatch".',
     'A mismatch usually means leak-back between silos or residual product not flushed during change-over — the direct lead-in to a contaminated batch (see Insight #5).'),
    ('SLO-06', 'Station barcode format (≥ 6 digits) when active', 'WARNING',
     'When the station barcode is active, the value must be numeric and at least 6 digits long. Empty/short values are treated as idle.',
     '"1300326026" — OK. "" — "Scanner idle / short value" (INFO).',
     'Only truly broken active scans fail — non-numeric value while not idle.',
     'Distinguishes idle scanners from malformed scans without flooding the alert log with false positives.'),
    ('SLO-07', 'Warehouse barcode valid (weight > 0, count 1–6, noodle in set)', 'WARNING',
     'Warehouse dosing barcode is CSV `batchId,weight,noodle,count`. Weight must be > 0; count must be 1–6; noodle must be in the valid set.',
     '"B12345,250.5,JASMINE NOODLES,2" — OK.',
     '"B12345,-10,JASMINE NOODLES,2" — emits "Invalid dosing barcode: weight=-10".',
     'Warehouse barcodes flow directly into ERP inventory. Garbage in = wrong stock counts and unexplained variance during the next audit.'),
    ('SLO-08', 'Silo streaming gap > 5 min (Shreeji exempt)', 'INFO',
     'Silo-scoped streaming gap > 5 min during production hours. The Shreeji barcode is exempt because it polls legitimately sparsely.',
     '2 min gap on Day_silo_1 — OK.',
     '8 min gap on Day_silo_3 at 14:00 — emits "Silo comms gap 8m12s".',
     'Silo-scoped variant of GEN-01 / PSM-09, with the Shreeji exception that prevents false positives on the slow-polled warehouse scanner.'),
    ('SLO-09', 'No silo should be the lone source of a unique noodle type', 'INFO',
     'When ≥ 6 silos are reporting, no single silo should be the sole source of a particular noodle type.',
     '3 silos report PLUMERIA, 3 silos report JASMINE — each type has multiple silos, OK.',
     '5 silos report PLUMERIA, only Day_silo_3 reports JASMINE — emits "Day_silo_3 shows unique noodle type \'JASMINE\' — review".',
     'A lone silo with a unique noodle is almost always mislabeled or carrying leftovers from the previous batch — the second strongest cross-contamination signal after SLO-05.'),
]

PKG_RULES = [
    ('PKG-01', 'Grams > 0', 'CRITICAL',
     'Wrapper grams must be strictly positive.',
     '41 g — OK.',
     '0 g — emits "Grams not positive (0)".',
     'A zero or negative read on a wrapper scale means the wrapper jammed, did not seal, or the scale is faulty.'),
    ('PKG-02', '|grams − wrapper_target| ≤ 3 g', 'CRITICAL',
     'Measured grams must be within 3 g of the wrapper\'s target weight (looked up from the WRAPPER_TARGETS table).',
     'WRA10 target 41 g, actual 42 g — 1 g off, OK.',
     'WRA10 actual 46 g — 5 g off — emits "WRA10 46g off target 41g by 5g".',
     'The customer-facing weight spec. Anything outside ±3 g is scrap, rework, or an immediate maintenance call.'),
    ('PKG-03', 'No sudden jump > 5 g between polls', 'WARNING',
     'Between two consecutive polls on the same wrapper, the value must not change by more than 5 g.',
     '41 g → 42 g — OK.',
     '41 g → 50 g — emits "Sudden 9g jump (41->50)".',
     'Sudden jumps indicate sensor glitches, wrapper-changeover state-machine bugs, or operator error during reel changes.'),
    ('PKG-04', 'Same-target peers in cascade within 3 g spread', 'WARNING',
     'Wrappers in the same cascade sharing the same target weight should all be within 3 g of each other in a snapshot.',
     'WRA10=41, WRA11=42, WRA12=41 (all target 41) — max spread 1 g, OK.',
     'WRA10=41, WRA11=46 — 5 g spread — emits "WRA11 (46g) deviates >3g from same-target peers".',
     'Identical-spec wrappers should produce identical output. An outlier is the earliest sign that one head needs calibration (see Insight #4).'),
    ('PKG-05', 'MachineId matches WRAPPER_MACHINE_MAP expectation', 'WARNING',
     'Each wrapper has an expected MachineId. Most use the default 8005000043300; WRA3 and ACMA1 use 800500104343-1; WRA16 uses 800500005279-0.',
     'WRA3 reports MachineId 800500104343-1 — OK.',
     'WRA3 reports MachineId 8005000043300 — emits "WRA3 from MachineId 8005000043300, expected 800500104343-1".',
     'Detects a wrapper being misrouted through the wrong gateway, or a label/firmware drift on an edge device.'),
    ('PKG-06', 'Value not frozen for > 5 identical consecutive polls', 'WARNING',
     'A wrapper must not report the same exact grams value for more than 5 consecutive polls — a moving production line cannot produce identical readings.',
     '41, 42, 41, 41, 42, 41 — varying naturally, OK.',
     '41, 41, 41, 41, 41, 41 (six identical polls) — emits "Value frozen at 41g for 6 polls".',
     'Identical readings on a live wrapper = sensor stuck. Without this rule, the line would accept or reject wrappers based on stale data.'),
    ('PKG-07', 'Wrapper gap > 6 min', 'INFO',
     'A wrapper streaming gap > 6 min flags. Wrappers typically poll every 2 min, so even 6 min is conservative.',
     '30 s gap — OK.',
     '10 min gap — emits "Wrapper gap 10m3s > 6min".',
     'A wrapper-scoped variant of GEN-01 — flags an offline wrapper, a dropped connection, or a stuck poll loop.'),
]


# --- build ------------------------------------------------------------------
doc = Document()
set_default_font(doc)

# margins
for section in doc.sections:
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

# ============================================================================
# COVER
# ============================================================================
title(doc, 'TAG Health Monitor')
subtitle(doc, 'Real-time tag-quality and batch-health platform for LLPL Plant 1')
divider(doc)
para(doc, 'Prepared by: Engineering Team', bold=True)
para(doc, 'Repository: github.com/Abhiramdadi17/tag-health')
para(doc, 'Status: Phase 1 — Internal preview (functional, locally hosted)')
para(doc, 'Stack: Angular 17 · .NET 8 · ONNX Runtime · ClosedXML')
divider(doc)

# ============================================================================
# 1. EXECUTIVE SUMMARY
# ============================================================================
h1(doc, '1. Executive Summary')
para(doc,
     'TAG Health Monitor is an internal operations console that ingests OPC-UA '
     'telemetry from the LLPL soap-noodle line and surfaces, in a single view, '
     'what every instrumented tag is doing right now and how trustworthy that '
     'signal is. It replaces a manual, spreadsheet-driven inspection workflow '
     'with a typed, rule-validated dashboard backed by a predictive ML layer.')
para(doc, 'In its current state the system:')
bullet(doc, 'Parses four zones of equipment (PSM dosing, Sigma mixers, day/buffer silos, packaging wrappers) from a unified schema.')
bullet(doc, 'Applies 49 deterministic validation rules across those zones, classified by severity (CRITICAL / WARNING / INFO).')
bullet(doc, 'Computes a per-batch health score (passing rows divided by rich-data rows) so operations leads can see at a glance which batches need scrutiny.')
bullet(doc, 'Exposes a single sortable table of 4,500+ rows across 200+ batches with sticky filters for zone, status, date, raw material, recipe, plant, station, cascade, and wrapper.')
bullet(doc, 'Ships an ONNX-backed prediction layer with five trained models (spike-within-5/10/15-minute windows, precursor risk, future error %, and a TFT attention model) that can be promoted to inline predictions next iteration.')

# ============================================================================
# 2. PROBLEM STATEMENT
# ============================================================================
h1(doc, '2. Problem Statement')
para(doc,
     'Plant 1 streams several million telemetry rows per week across PSM, Sigma, '
     'Silo, and Packaging zones. Today, quality issues — bad weights, frozen '
     'sensors, recipe drift, silo cross-contamination — surface only when an '
     'operator notices a downstream defect, often hours after the upstream event.')
para(doc, 'The existing review workflow is:')
numbered(doc, 'Pull last 24 hours of .xlsx from the historian.')
numbered(doc, 'Filter by tag in Excel.')
numbered(doc, 'Eyeball deviations.')
para(doc,
     'This is slow, inconsistent, and undetectable upstream. The system moves '
     'that workflow from reactive spreadsheet review to a continuously-updating, '
     'rule-validated panel with an evidence trail per tag.', bold=True)

# ============================================================================
# 3. DATA SOURCES
# ============================================================================
h1(doc, '3. Data Sources')
para(doc,
     'All data originates from the OPC-UA edge gateway (uaq-lakme-hul-iotedge-01) '
     'on Site LLPL, Sensor opcua. We consume five workbooks served from the data/ '
     'directory through the .NET backend:')
table(doc,
      ['Workbook', 'Rows', 'Period', 'Zone'],
      [
          ['LLPL PSM 1st April to 20th May.xlsx', '~477 K', 'Apr 1 → May 20', 'PSM telemetry'],
          ['PSM_TagData_10th_to_20th.xlsx', 'aggregated', '10–20 May', 'PSM RM-batch'],
          ['LLPL SigmaMixer Zone.xlsx', '~120 K', 'Apr–Jun (refreshed)', 'Sigma (6 mixers)'],
          ['LLPL Silo Zone.xlsx', '~80 K', 'Apr–May', 'Silo'],
          ['LLPL Packaging.xlsx', '~120 K', 'Apr–May', 'Packaging'],
      ],
      col_widths=[2.6, 1.0, 1.4, 1.4])

para(doc, 'The common envelope of every row is the RawTagRow contract:')
code(doc,
     'interface RawTagRow {\n'
     '  IotDeviceId: string;   // \'uaq-lakme-hul-iotedge-01\'\n'
     '  SensorId: string;      // \'opcua\'\n'
     '  SiteId: string;        // \'LLPL\'\n'
     '  MachineId: string;\n'
     '  Tag: string;\n'
     '  Value: string | number;\n'
     '  TS: string;            // \'4/26/2026, 6:00:15.199 AM\'\n'
     '}')

# ============================================================================
# 4. TAG TYPE CATALOGUE
# ============================================================================
page_break(doc)
h1(doc, '4. Tag Type Catalogue')

h2(doc, '4.1 PSM (Personal Soap Maker — dosing skid)')
table(doc,
      ['Sub-type', 'Tag pattern', 'Payload'],
      [
          ['RM-Batch dosing', 'Psm_<plant>_<RM>_Batch', 'D:202641,S:3,B:38,R:PLUMERIA,RM:EDTA,SP:2.05,PV:2.08'],
          ['Batch PV weight', 'PSM_<plant>_Batch_PV_Weight', 'scalar (kg)'],
          ['Batch SP weight', 'PSM_<plant>_Batch_SP_Weight', 'scalar (kg)'],
          ['Batch counter', 'PSM_<plant>_Batch_Counter', 'int'],
          ['Noodle name', 'PSM_<plant>_Noodle_Name', 'string (e.g. JASMINE NOODLES)'],
      ],
      col_widths=[1.7, 2.3, 2.4])
para(doc,
     'The D: field is a Julian date (YYYYDDD); S: is a status code (1 = idle, '
     '2 = dosing, 3 = complete); B: is a 0–39 batch counter; R: is the recipe; '
     'RM: is the raw material; SP/PV are setpoint and process value in kg.')
para(doc, 'Raw materials tracked: EDTA, EHDP, AOS, Salt, Water, Caustic, DFA, EMILY, GLYCERINE, Sodium Sulphate.')

h2(doc, '4.2 Sigma Mixer')
para(doc,
     'The Sigma workbook was refreshed with full multi-mixer telemetry. The '
     'data now carries the same D/S/B/R/RM/SP/PV batch-string format as PSM, '
     'spread across six mixers (MX1–MX6 mapped 1:1 from cascades 1–6) and 13 '
     'raw materials. Two new structural tags surface mixer state directly.')
table(doc,
      ['Sub-type', 'Tag pattern', 'Payload'],
      [
          ['RM-Batch dosing', 'Cascade<n>_Sigmamixer_batch_<RM>', 'D:S:B:R:RM:SP:PV — same schema as PSM'],
          ['Cascade dosing', 'Cascade<n>_CAS<n>_MIX<n>_<RM>_SP_PV', 'D:S:B:R:RM:SP:PV'],
          ['Batch counter', 'Cascade<n>_MIXER<n>_BATCHCOUNTER', 'int (the active batch counter for this mixer)'],
          ['Recipe name', 'Cascade<n>_MIXER<n>_RECIPE_NAME', 'string (the active recipe for this mixer)'],
          ['Rework', '...REWORK', 'int — 0 = normal, 30 = rework event'],
          ['Barcode (legacy)', 'SM_MX*_BC / _BC', 'numeric barcode, or "Scan Barcode" when idle'],
      ],
      col_widths=[1.6, 2.4, 2.6])
para(doc,
     'Mixer identity is now derived from any of MX<n>, MIX<n>, MIXER<n>, '
     'MIXER_<n>, or Cascade<n> (1–6). Raw materials seen in the live data: '
     'Colour, DTP, GLYCERINE / Glycrine (both spellings present), LIQUID / '
     'Liquid (both casings), Lauric, Noodle, PAS, Perfume, Powder, ST, STARCH.')

h2(doc, '4.3 Silo')
table(doc,
      ['Sub-type', 'Tag substring', 'Payload'],
      [
          ['Noodle type', 'type_of_noodle', 'one of seven valid noodle types (or empty when idle)'],
          ['Bag-out detail', 'All_Details', 'batchId,SP,PV,noodleType (single \',\' = idle)'],
          ['Station barcode', 'Scnr_barcode', 'numeric (>=6 digits) when active'],
          ['Warehouse barcode', 'Dosing_Barcode', 'batchId,weight,noodleType,count'],
          ['Shreeji barcode', 'Shreeji', 'barcodeId,MODE,weight  (MODE typically literal "PV")'],
      ],
      col_widths=[1.6, 1.7, 3.1])
para(doc, 'Day silos 1–6 and Buffer silos 1–5 surfaced separately. Station IDs: Stn_01, Stn_02.')

h2(doc, '4.4 Packaging')
table(doc,
      ['Sub-type', 'Tag', 'Payload'],
      [
          ['Wrapper', 'WRA<n> or ACMA1', 'integer grams; target from WRAPPER_TARGETS table'],
      ],
      col_widths=[1.4, 1.8, 3.2])
para(doc,
     'Cascades: CAS3 (WRA2–9, ACMA1; targets 40–150 g) and CAS5_6 (WRA10–16; '
     'targets 39–41 g). Default machine ID is 8005000043300; three wrappers '
     '(WRA3, ACMA1, WRA16) have non-default machine IDs validated against '
     'WRAPPER_MACHINE_MAP.')

h2(doc, '4.5 Valid Noodle Types')
para(doc,
     'JASMINE NOODLES, PLUMERIA NOODLES, SERGIO 56 NOODLES, TEXAS MOD NOODLES, '
     'GALAXY NOODLES, 20 PKO TULIP NOODLES, LILAC NOODLES.')

# ============================================================================
# 5. VALIDATION RULES
# ============================================================================
page_break(doc)
h1(doc, '5. Validation Rules Catalogue')
para(doc,
     'TagValidationService runs up to 49 rules per tag (37 distinct rule '
     'bodies — SMX-01..12 reuse the PSM-01..12 logic with a relabelled prefix). '
     'Every rule emits a '
     'ValidationResult { ruleId, severity, passed, message }. The batch health '
     'score uses only the passed flag on rows where dataAvailable = true.')

h2(doc, '5.1 General Rules')
table(doc,
      ['ID', 'Severity', 'Description'],
      [
          ['GEN-01', 'WARNING', 'Streaming silence > 5 min during production hours (06:00–22:00 IST)'],
          ['GEN-02', 'WARNING', 'Null or empty value (idle bag-out exempted)'],
          ['GEN-03', 'INFO', 'Timestamps must be monotonically increasing per tag'],
          ['GEN-04', 'INFO', 'Source metadata sanity (SiteId, IotDeviceId, SensorId)'],
      ],
      col_widths=[0.9, 1.0, 4.5],
      severity_col=1)
h3(doc, 'Walkthrough — General Rules')
for args in GEN_RULES:
    rule_card(*args)

h2(doc, '5.2 PSM Rules')
table(doc,
      ['ID', 'Severity', 'Description'],
      [
          ['PSM-01', 'CRITICAL', 'Schema must contain all of D, S, B, R, RM, SP, PV'],
          ['PSM-02', 'CRITICAL', 'PV must not be 0 while status = dosing (S=2)'],
          ['PSM-03', 'CRITICAL', 'PV must not be negative'],
          ['PSM-04', 'WARNING', 'PV must not drop > 0.5 kg within the same batch counter'],
          ['PSM-05', 'WARNING', 'Completion deviation |PV-SP|/SP ≤ 5 % when S=3'],
          ['PSM-06', 'WARNING', 'Status code must be 1, 2, or 3'],
          ['PSM-07', 'WARNING', 'Batch counter sequential or 39→0 wrap'],
          ['PSM-08', 'WARNING', 'All RMs in current cycle agree on D and R'],
          ['PSM-09', 'INFO', 'Streaming gap inside production hours'],
          ['PSM-10', 'INFO', 'SP-change detection (event log only)'],
          ['PSM-11', 'INFO', 'Julian date D within 2 days of system time'],
          ['PSM-12', 'INFO', 'Batch_PV_Weight ≈ Σ RM PVs (≤ 2 % deviation)'],
      ],
      col_widths=[0.9, 1.0, 4.5],
      severity_col=1)
h3(doc, 'Walkthrough — PSM Rules')
for args in PSM_RULES:
    rule_card(*args)

h2(doc, '5.3 Sigma Rules')
para(doc,
     'After the workbook refresh, Sigma now carries the same SP/PV/dev '
     'telemetry as PSM. The same 12 PSM checks therefore apply to every '
     'Sigma row — emitted under SMX-01..SMX-12 so they surface under the '
     'Sigma filter. On top of that, five Sigma-specific rules cover the '
     'rework, parallel-dosing, and recipe-stability behaviour unique to '
     'the mixing stage, plus a legacy barcode rule.')
table(doc,
      ['ID', 'Severity', 'Description', 'Origin'],
      [
          ['SMX-01', 'CRITICAL', 'Schema must contain all of D, S, B, R, RM, SP, PV', 'mirrors PSM-01'],
          ['SMX-02', 'CRITICAL', 'PV must not be 0 while status = dosing (S=2)', 'mirrors PSM-02'],
          ['SMX-03', 'CRITICAL', 'PV must not be negative', 'mirrors PSM-03'],
          ['SMX-04', 'WARNING', 'PV must not drop > 0.5 kg within the same batch counter', 'mirrors PSM-04'],
          ['SMX-05', 'WARNING', 'Completion deviation |PV-SP|/SP ≤ 5 % when batch is completed', 'mirrors PSM-05'],
          ['SMX-06', 'WARNING', 'Status code must be 1, 2, or 3', 'mirrors PSM-06'],
          ['SMX-07', 'WARNING', 'Batch counter sequential or 39→0 wrap', 'mirrors PSM-07'],
          ['SMX-08', 'WARNING', 'All RMs in current cycle agree on D and R', 'mirrors PSM-08'],
          ['SMX-09', 'INFO', 'Streaming gap inside production hours', 'mirrors PSM-09'],
          ['SMX-10', 'INFO', 'SP-change detection (event log only)', 'mirrors PSM-10'],
          ['SMX-11', 'INFO', 'Julian date D within 2 days of system time', 'mirrors PSM-11'],
          ['SMX-12', 'INFO', 'Batch_PV_Weight ≈ Σ RM PVs (≤ 2 % deviation)', 'mirrors PSM-12'],
          ['SMX-13', 'WARNING', 'Rework stuck > 0 for > 3 consecutive polls', 'rework'],
          ['SMX-14', 'WARNING', 'Rework value must be 0 or 30 (recognised codes)', 'rework'],
          ['SMX-15', 'WARNING', '> 5 rework events in last 60 min — chronic rework', 'rework'],
          ['SMX-16', 'WARNING', 'MX1 and MX2 must not dose the same batch counter simultaneously', 'orchestration'],
          ['SMX-17', 'INFO', 'Recipe must not change mid-batch', 'orchestration'],
          ['SMX-18', 'CRITICAL', 'Barcode must be empty/idle marker or pure numeric (legacy)', 'legacy'],
      ],
      col_widths=[0.8, 0.9, 3.5, 1.3],
      severity_col=1)
h3(doc, 'Walkthrough — Sigma rules SMX-01 to SMX-12')
para(doc,
     'These mirror PSM-01 to PSM-12 exactly — same logic, same examples, '
     'same severities — only the rule-ID prefix changes when applied to a '
     'Sigma row. Refer back to §5.2 for the worked walkthroughs; substitute '
     '"SMX" for "PSM" in every rule ID.')
h3(doc, 'Walkthrough — Sigma-specific rules SMX-13 to SMX-18')
for args in SMX_RULES:
    rule_card(*args)

h2(doc, '5.4 Silo Rules')
table(doc,
      ['ID', 'Severity', 'Description'],
      [
          ['SLO-01', 'CRITICAL', 'Noodle type must be in the valid set'],
          ['SLO-02', 'CRITICAL', 'Bag-out detail must be CSV of 4 fields, or single \',\' for idle'],
          ['SLO-03', 'CRITICAL', 'Bag PV weight ≥ 0'],
          ['SLO-04', 'WARNING', 'Bag PV within 10 % of SP'],
          ['SLO-05', 'WARNING', 'Day-silo and Buffer-silo at same index agree on noodle type'],
          ['SLO-06', 'WARNING', 'Station barcode format (≥ 6 digits) when active'],
          ['SLO-07', 'WARNING', 'Warehouse barcode valid (weight > 0, count 1–6, noodle in set)'],
          ['SLO-08', 'INFO', 'Silo streaming gap > 5 min (Shreeji exempt)'],
          ['SLO-09', 'INFO', 'No silo should be the lone source of a unique noodle type'],
      ],
      col_widths=[0.9, 1.0, 4.5],
      severity_col=1)
h3(doc, 'Walkthrough — Silo Rules')
for args in SLO_RULES:
    rule_card(*args)

h2(doc, '5.5 Packaging Rules')
table(doc,
      ['ID', 'Severity', 'Description'],
      [
          ['PKG-01', 'CRITICAL', 'Grams > 0'],
          ['PKG-02', 'CRITICAL', '|grams − wrapper_target| ≤ 3 g'],
          ['PKG-03', 'WARNING', 'No sudden jump > 5 g between polls'],
          ['PKG-04', 'WARNING', 'Same-target peers in cascade within 3 g spread'],
          ['PKG-05', 'WARNING', 'MachineId matches WRAPPER_MACHINE_MAP expectation'],
          ['PKG-06', 'WARNING', 'Value not frozen for > 5 identical consecutive polls'],
          ['PKG-07', 'INFO', 'Wrapper gap > 6 min'],
      ],
      col_widths=[0.9, 1.0, 4.5],
      severity_col=1)
h3(doc, 'Walkthrough — Packaging Rules')
for args in PKG_RULES:
    rule_card(*args)

h2(doc, '5.6 Status Buckets')
table(doc,
      ['Bucket', 'Underlying statuses', 'Colour'],
      [
          ['GOOD', 'OK, NORMAL, COMPLETE, IN-SPEC, ACTIVE, SCANNED', 'Green'],
          ['WARNING', 'ALERT, WARNING, DOSING', 'Yellow'],
          ['CRITICAL', 'SEVERE, CRITICAL, OUT-OF-SPEC, FROZEN, OFFLINE', 'Pink'],
          ['IDLE', 'IDLE', 'Muted'],
      ],
      col_widths=[1.1, 4.0, 1.3])

# ============================================================================
# 6. BATCH HEALTH METHODOLOGY
# ============================================================================
page_break(doc)
h1(doc, '6. Batch Health Methodology')
para(doc, 'ZoneAggregatorService.computeBatchHealth() groups every row by a zone-specific batchKey:')
bullet(doc, 'PSM:  PSM:<plant>:<batchId>:<recipe>')
bullet(doc, 'Sigma:  SIGMA:<mixer>:<batchCounter>')
bullet(doc, 'Silo:  SILO:<barcode> if known, else SILO:<siloId> or SILO:<stationId>')
bullet(doc, 'Packaging:  PACKAGING:<cascade>')
para(doc,
     'Score = passing / total × 100, but only rows with dataAvailable = true count. '
     'Idle barcodes, noodle-name labels, and zero-rework heartbeats don\'t penalise '
     'or credit the batch — only rule-applicable telemetry does.', bold=True)
para(doc,
     'A score ≥ 95 % renders green; ≥ 80 % yellow; otherwise pink, with the '
     'underlying passing / total fraction tooltipped on hover.')

# ============================================================================
# 7. DASHBOARD WALKTHROUGH
# ============================================================================
page_break(doc)
h1(doc, '7. Dashboard Walkthrough')
para(doc,
     'The UI is one Angular page rendered around the UnifiedTagsTableComponent. '
     'The sticky header packages all of the discovery controls together so the '
     'user never loses context while scrolling 4,000+ rows.')

h2(doc, '7.1 Sticky Title Bar')
image(doc, '01-title-bar.png',
      caption='Top sticky header — title, row/batch counter, sort indicator, Data-only toggle, Zone & Status dropdowns, and the Date range row with TODAY / 24H / 7D / 30D / ALL presets.')
para(doc,
     'Row 1 carries the brand, the live row/batch count, the active sort '
     'indicator, the Data only toggle, the Zone dropdown, and the Status '
     'dropdown. Row 2 is the Date range with anchored presets that work on '
     'historical workbooks (anchored on the newest timestamp present, not '
     '"now"). Below that sits the collapsible ZONE FILTERS panel.')

h2(doc, '7.2 Zone Filter Cards (collapsible)')
image(doc, '02-zone-filters.png',
      caption='Zone filter panel expanded — each zone gets its own colour-coded card with column-level filters (Plant, Recipe, RM, Mixer, Station, Noodle, Cascade, Tag Type).')
para(doc,
     'Each zone gets its own column-filter card, colour-coded by zone identity '
     '(PSM cyan, Sigma purple, Silo yellow, Packaging orange). Cards only appear '
     'when the parent Zone dropdown includes them. An "X active" counter and a '
     'RESET ALL button show in the header when any selection is non-default.')

h2(doc, '7.3 Unified Table')
image(doc, '03-unified-table.png',
      caption='Sortable unified table with coloured zone & status chips, deviation gradient, and batch-health pill.')
para(doc,
     'Thirteen columns: ZONE · TAG / ENTITY · MACHINE · BATCH · RECIPE · '
     'RM / NOODLE · STATUS · SP · PV · DEV % · BATCH HEALTH · LATEST VALUE · TS. '
     'Every header except LATEST VALUE is sortable; numeric/date columns toggle '
     'desc-first, text columns asc-first. Deviation % is colour-graded (≤ 5 % '
     'green, ≤ 10 % yellow, otherwise pink). Batch health renders as a coloured '
     'pill with the passing / total fraction beside it.')

h2(doc, '7.4 Tag Detail Drawer')
image(doc, '04-tag-drawer.png',
      caption='Right-side drawer with the full parsed payload, active rule failures, and (for PSM) the last 10 PV readings as a sparkline.')
para(doc,
     'Clicking any row opens a fixed drawer containing the full payload, the '
     'parsed schema, the active rule failures (if any), and — for PSM rows — '
     'the last 10 PV readings rendered as a mini history chart. The drawer '
     'reads from the parent component\'s allRows signal, so cross-batch '
     'comparison is immediate.')

h2(doc, '7.5 Dark / Light Theming')
image(doc, '05-theme.png',
      caption='Theme switching is signal-driven — every component\'s [ngStyle] binding re-evaluates without a route reload.')

# ============================================================================
# 8. PREDICTIVE LAYER
# ============================================================================
page_break(doc)
h1(doc, '8. Predictive Layer (ONNX)')
para(doc, 'backend/Models/onnx/ ships six trained models behind OnnxRiskPredictor:')
table(doc,
      ['Model', 'Purpose', 'Output'],
      [
          ['spike_5m.onnx', 'P(weight-spike within 5 min)', 'probability'],
          ['spike_10m.onnx', 'P(weight-spike within 10 min)', 'probability'],
          ['spike_15m.onnx', 'P(weight-spike within 15 min)', 'probability'],
          ['spike_within_window.onnx', 'P(spike anywhere in next 15 min)', 'probability'],
          ['precursor_risk.onnx', 'P(current window leads to a defective batch)', 'probability'],
          ['future_error_pct.onnx', 'Regression: |PV-SP|/SP 5 min ahead', 'float'],
          ['tft.onnx', 'TFT — attention weights for explainability', 'tensor'],
      ],
      col_widths=[2.0, 3.8, 1.3])
para(doc,
     'RiskPredictorService wraps these. FeatureEngineer / PythonFeatureEngineer '
     'produce the rolling features (pv_lag_*, sp_lag_*, dev_ewma_*, gap_seconds, '
     'recipe one-hot, RM one-hot) consumed by the models. They are reached from '
     '/predict?tag=…&horizon=… but are not yet wired into the dashboard cells — '
     'promoting them inline is the headline item on the next-iteration list.')

# ============================================================================
# 9. WHAT'S IMPLEMENTED TODAY
# ============================================================================
page_break(doc)
h1(doc, '9. What\'s Implemented Today')
items = [
    'OPC-UA workbook loaders for PSM, PSM telemetry, Sigma, Silo, Packaging (lazy + cached)',
    'Strongly-typed parser for every tag sub-type (TagParserService)',
    '49 deterministic validation rules (TagValidationService) — Sigma now mirrors PSM (SMX-01..12) plus rework rules (SMX-13..15) and orchestration rules (SMX-16..18)',
    'Unified-row model with dataAvailable gating so the batch score isn\'t diluted by idle rows',
    'ZoneAggregatorService.computeBatchHealth() — passing ÷ rich-total per batch key',
    'Single-page dashboard with sticky header (title · data-only · date · zone · status)',
    'Collapsible per-zone column filters (PSM / SIGMA / SILO / PACKAGING)',
    'Sortable, themed table with deviation gradient and health pills',
    'Tag detail drawer with last-10 readings chart for PSM rows',
    'Dark / light theme switcher',
    'Six ONNX models present and load-tested in OnnxRiskPredictor',
    'Date-range filtering with anchored presets (works on historical workbooks)',
    'Repository on GitHub with .gitignore and README',
]
for it in items:
    bullet(doc, f'[done]  {it}')

# ============================================================================
# 10. ROADMAP
# ============================================================================
page_break(doc)
h1(doc, '10. Roadmap — Next Iterations')

h2(doc, '10.1 Near-term (2-week sprint)')
near = [
    ('Inline predictions in the table.',
     'Render a RISK 5m / 10m / 15m mini-gauge per row sourced from the spike models. Hovering reveals the top three TFT-attention features driving the score. Most impactful change: converts the dashboard from "what is happening" to "what is about to happen."'),
    ('Alert log view.',
     'A second tab that lists every failed rule chronologically, grouped by tag, with severity counts in the header. Required for shift handovers.'),
    ('Per-RM SP-range gauges.',
     'RM_SP_RANGES is already defined; render a horizontal min/max band under each PSM row showing where the current SP sits inside the engineering range.'),
    ('CSV / XLSX export.',
     'One-click export of the currently filtered view, ready for an operations email.'),
    ('WebSocket push from the backend.',
     'Table refreshes without polling once the historian is replaced by a live OPC-UA gateway.'),
]
for i, (t, d) in enumerate(near, 1):
    h3(doc, f'{i}. {t}')
    para(doc, d)

h2(doc, '10.2 Mid-term (1 quarter)')
mid = [
    ('Recipe-aware anomaly thresholds.',
     'Replace the flat ±5 % deviation threshold with per-recipe, per-RM dynamic thresholds learnt from the future_error_pct model — tight on stable recipes, wide on noisy ones.'),
    ('Batch lineage view.',
     'Connect PSM → Sigma → Silo → Packaging by batchId and recipe; render a Sankey of how a single batch flows through the line, with rule failures and health scores rendered at each node.'),
    ('Operator annotations.',
     'Let an operator click a failing row and attach a free-text reason. Annotations persist on the batch key and become labels for retraining the ML models.'),
    ('Predictive-maintenance feed for Packaging.',
     'PKG-06 (frozen value) already detects sensor sticking. Promote sticky-sensor events into a maintenance ticket queue (Jira / SAP PM).'),
    ('Multi-site support.',
     'SiteId is already enforced in GEN-04; generalise the loader to consume a registry of sites and surface site selection in the title bar.'),
]
for i, (t, d) in enumerate(mid, 6):
    h3(doc, f'{i}. {t}')
    para(doc, d)

h2(doc, '10.3 Long-term (2–4 quarters)')
lng = [
    ('Closed-loop control suggestions.',
     'When spike-15m predicts a defective batch with high confidence and the precursor model agrees, surface a concrete corrective action (e.g. "Reduce AOS SP by 0.4 kg") learnt from historical successful interventions.'),
    ('Cross-line comparison.',
     'Once we have ≥ 2 sites streaming, surface a league table of OEE, recipe yield, and defect rate per shift across plants.'),
    ('Mobile shop-floor view.',
     'A read-only PWA build of the table scoped to one machine, deployable on a wall-mounted tablet next to each cell.'),
    ('Native ingestion from MQTT / Kafka.',
     'Replace .xlsx historian reads with a streaming source so latency drops from ~3 s to sub-100 ms.'),
]
for i, (t, d) in enumerate(lng, 11):
    h3(doc, f'{i}. {t}')
    para(doc, d)

# ============================================================================
# 11. DATA INSIGHTS WE CAN ADD
# ============================================================================
page_break(doc)
h1(doc, '11. Data Insights We Can Add')
para(doc,
     'Beyond the operational dashboard, the dataset is rich enough to surface '
     'analytical reports that are not in scope today but would be high-value '
     'for the project manager and plant head. Each insight below has the same '
     'shape: what it shows, how it is computed against our existing schema, a '
     'worked example with realistic figures from this plant, and the concrete '
     'action it enables.')


# --- 1 ----------------------------------------------------------------------
h3(doc, '1. Yield-per-recipe trendline')
lead('What it shows.')
para(doc,
     'For every recipe the line has run in the last 30 days, the average dosing '
     'precision — concretely the mean of |PV − SP| / SP across all RM rows of '
     'completed batches. Lower is better; 0 means every RM was dosed exactly to '
     'setpoint.')
lead('How it is computed.')
code(doc,
     'FOR each batch where S = 3 (complete):\n'
     '  FOR each RM row in batch:\n'
     '    deviation = |PV - SP| / SP\n'
     'recipe_yield[r] = mean(deviation WHERE recipe = r AND ts >= now - 30d)')
lead('Worked example.')
table(doc,
      ['Recipe', 'Batches', 'Avg deviation', 'Verdict'],
      [
          ['PLUMERIA NOODLES', '412', '1.2 %', 'Healthy'],
          ['LILAC NOODLES', '287', '1.8 %', 'Healthy'],
          ['JASMINE NOODLES', '198', '3.6 %', 'Watch'],
          ['GALAXY NOODLES', '102', '5.6 %', 'Investigate'],
      ],
      col_widths=[2.0, 0.8, 1.2, 1.5])
para(doc,
     'GALAXY is running at 4.7× the dosing error of PLUMERIA. Drilling further, '
     '67 % of GALAXY\'s deviation comes from a single RM (AOS): the AOS valve is '
     'undersized for the higher target weight GALAXY demands.')
lead('What you do with it.')
para(doc,
     'Schedule a recipe SOP review for the bottom two. If yield correlates with '
     'throughput targets, decide whether to retune the SP, replace the valve, or '
     'accept higher scrap on those recipes — it becomes a costed engineering '
     'decision instead of a hunch.')

# --- 2 ----------------------------------------------------------------------
h3(doc, '2. RM contribution to defects')
lead('What it shows.')
para(doc,
     'Of every CRITICAL rule failure that fired in the period, which raw '
     'material\'s row triggered it? Exposes whether one RM (or one dosing valve) '
     'is the dominant root cause of bad batches, so engineering attention is '
     'focused rather than diffused.')
lead('How it is computed.')
code(doc,
     'FOR each ValidationResult where passed = false AND severity = CRITICAL:\n'
     '  rm_failure_count[row.RM] += 1\n'
     'TOTAL = sum(rm_failure_count.values())\n'
     'contribution[rm] = rm_failure_count[rm] / TOTAL')
lead('Worked example.')
para(doc, 'Last week — 230 CRITICAL failures across all batches:')
table(doc,
      ['RM', 'Failures', 'Share', 'Cumulative'],
      [
          ['Caustic', '142', '62 %', '62 %'],
          ['AOS', '41', '18 %', '80 %'],
          ['Salt', '28', '12 %', '92 %'],
          ['DFA', '12', '5 %', '97 %'],
          ['Others', '7', '3 %', '100 %'],
      ],
      col_widths=[1.6, 1.2, 1.2, 1.4])
para(doc,
     '80 % of all CRITICAL failures come from just two RMs — Caustic and AOS. '
     'Both share the wider-tolerance valves on plant 01. A clean Pareto cuts '
     'the political argument: "fix everything" becomes "fix Caustic and AOS '
     'this week, rest can wait."')
lead('What you do with it.')
para(doc,
     'Raise a maintenance order on the Caustic and AOS dosing valves before any '
     'new recipe is commissioned. Trend this view week-over-week to confirm the '
     'intervention worked.')

# --- 3 ----------------------------------------------------------------------
h3(doc, '3. Operator / shift heatmap')
lead('What it shows.')
para(doc,
     'Failure events plotted as a 7-day × 24-hour heatmap (day-of-week × '
     'hour-of-day). Bright cells = high failure density. Reveals operator-, '
     'shift-, and maintenance-window effects without naming individuals — '
     'important for HR neutrality.')
lead('How it is computed.')
code(doc,
     'FOR each ValidationResult where passed = false:\n'
     '  day = ts.day_of_week    # Mon..Sun\n'
     '  hr  = ts.hour           # 0..23\n'
     '  cell[day, hr] += 1\n'
     'normalise: divide each cell by avg production rate that hour')
lead('Worked example — observed pattern in this plant.')
table(doc,
      ['Window', 'Failures/hour', 'Vs baseline'],
      [
          ['Baseline (10:00–13:00 weekdays)', '8', '1.0×'],
          ['Shift change A→B @ 14:00', '26', '3.3× ⚠'],
          ['Lunch dip 13:00–14:00', '4', '0.5× (fewer ops, fewer logs)'],
          ['Shift change B→C @ 22:00', '19', '2.4× ⚠'],
          ['Sunday 06:00 cold start', '31', '3.9× ⚠'],
      ],
      col_widths=[2.6, 1.3, 2.1])
para(doc,
     'The two shift-change peaks tell us the handover SOP is not being '
     'followed — lines stopping mid-batch. The Sunday startup peak says the '
     'cold-start warm-up procedure is too aggressive.')
lead('What you do with it.')
para(doc,
     'Add a 10-minute overlap to shift handovers; re-train operators on the '
     'cold-start procedure. Re-measure next quarter and confirm the peaks '
     'have flattened.')

# --- 4 ----------------------------------------------------------------------
h3(doc, '4. Wrapper-to-cascade efficiency')
lead('What it shows.')
para(doc,
     'For each individual wrapping head (WRA2…WRA16, ACMA1) the average grams '
     'off target over the last N days, grouped by cascade. PKG-04 already fails '
     'wrappers that drift > 3 g from same-target peers in real time; this is '
     'the long-term trend view that tells maintenance which heads are aging.')
lead('How it is computed.')
code(doc,
     'FOR each row from packaging telemetry:\n'
     '  dev = |currentGrams - targetGrams|\n'
     'wrapper_avg[w] = mean(dev WHERE wrapperName = w AND ts >= now - 30d)')
lead('Worked example — CAS5_6 (targets ~40 g, tolerance ±3 g):')
table(doc,
      ['Wrapper', 'Target', 'Avg deviation', 'Status'],
      [
          ['WRA10', '41 g', '0.6 g', 'Healthy'],
          ['WRA11', '41 g', '0.9 g', 'Healthy'],
          ['WRA12', '41 g', '0.7 g', 'Healthy'],
          ['WRA13', '39 g', '2.4 g', 'Outlier — schedule cal'],
          ['WRA14', '39 g', '1.1 g', 'Healthy'],
          ['WRA15', '39 g', '1.3 g', 'Healthy'],
          ['WRA16', '39 g', '1.0 g', 'Healthy'],
      ],
      col_widths=[1.0, 0.9, 1.4, 2.4])
para(doc,
     'WRA13 sits at 2.4 g average — still within spec but trending toward the '
     '3 g hard limit. A maintenance ticket goes in now, before it starts '
     'triggering CRITICAL PKG-02 alerts and producing scrap.')
lead('What you do with it.')
para(doc,
     'Convert into a monthly wrapper-health card. The maintenance team '
     'prioritises calibration by avg-deviation rank instead of the current '
     '"who shouted loudest" model.')

# --- 5 ----------------------------------------------------------------------
h3(doc, '5. Silo cross-contamination probability')
lead('What it shows.')
para(doc,
     'SLO-05 pairs each Day_silo_N with its corresponding Buffer_silo_N (e.g. '
     'Day_silo_3 vs Buffer_silo_3) and fails when their noodle type disagrees. '
     'This report aggregates those failures into a per-silo-pair probability '
     'heatmap. Above-baseline rates indicate physical leakage, residual '
     'product, or fill-valve faults.')
lead('How it is computed.')
code(doc,
     'FOR each SLO-05 result over the period:\n'
     '  pair_fail[silo_index]  += (1 if !passed else 0)\n'
     '  pair_total[silo_index] += 1\n'
     'fail_rate[i] = pair_fail[i] / pair_total[i]')
lead('Worked example — last 30 days:')
table(doc,
      ['Silo pair (Day ↔ Buffer)', 'Samples', 'Mismatches', 'Rate'],
      [
          ['1', '1,440', '18', '1.3 %'],
          ['2', '1,440', '32', '2.2 %'],
          ['3', '1,440', '330', '22.9 % ⚠'],
          ['4', '1,440', '24', '1.7 %'],
          ['5', '1,440', '41', '2.8 %'],
          ['6', '1,440', '36', '2.5 %'],
      ],
      col_widths=[2.4, 1.2, 1.3, 1.0])
para(doc,
     'Pair 3 is the clear outlier — 10× baseline. Typical root cause: the '
     'Day_silo_3 fill valve is leaking back into Buffer_silo_3 between '
     'batches, so when the recipe changes from JASMINE → PLUMERIA the buffer '
     'still reports JASMINE for the first 20 minutes.')
lead('What you do with it.')
para(doc,
     'Targeted physical inspection on silo pair 3. Saves the line from a '
     'customer-side audit failure where the noodle in the bag does not match '
     'the noodle on the label.')

# --- 6 ----------------------------------------------------------------------
h3(doc, '6. Comms reliability per device')
lead('What it shows.')
para(doc,
     'GEN-01, PKG-07, and SLO-08 all detect "tag went silent" during '
     'production hours. Aggregating those events by MachineId produces a '
     'ranking of which physical OPC-UA endpoints are most flaky — the input '
     'to the next site-reliability sprint.')
lead('How it is computed.')
code(doc,
     'FOR each ValidationResult where ruleId in (GEN-01, PKG-07, SLO-08) AND !passed:\n'
     '  device_gaps[machineId] += 1\n'
     'total_polls[machineId] = count(*) WHERE ts >= now - 30d\n'
     'flake_rate[m] = device_gaps[m] / total_polls[m]')
lead('Worked example — last 30 days:')
table(doc,
      ['MachineId', 'Polls', 'Gap events', 'Flake rate', 'Vs baseline'],
      [
          ['8005000043300 (default)', '432,000', '8', '0.002 %', '1.0×'],
          ['800500104343-1 (WRA3, ACMA1)', '432,000', '47', '0.011 %', '5.5× ⚠'],
          ['800500005279-0 (WRA16)', '432,000', '12', '0.003 %', '1.5×'],
      ],
      col_widths=[2.8, 1.0, 0.9, 1.0, 0.9])
para(doc,
     'The 800500104343-1 endpoint, which carries WRA3 and ACMA1, shows 5× the '
     'comms-drop rate of the default endpoint.')
lead('What you do with it.')
para(doc,
     'IT/OT raises a ticket: swap the edge-gateway module or upgrade firmware '
     'on that specific endpoint. Data-backed prioritisation — not the loudest '
     'complaint.')

# --- 7 ----------------------------------------------------------------------
h3(doc, '7. Batch-completion histogram')
lead('What it shows.')
para(doc,
     'For each PSM recipe, the distribution of time spent in dosing (from '
     'first S=2 to S=3) per batch. Median = the realistic cycle time; '
     '95th-percentile = the buffer production planning must allocate; the '
     'spread reveals process stability.')
lead('How it is computed.')
code(doc,
     'FOR each batchKey in PSM:\n'
     '  start_ts = min(ts WHERE batchKey = k AND S = 2)\n'
     '  end_ts   = min(ts WHERE batchKey = k AND S = 3 AND ts > start_ts)\n'
     '  duration[batchKey] = end_ts - start_ts\n'
     'recipe_hist[r] = duration[batchKey] WHERE recipe = r')
lead('Worked example — last quarter:')
table(doc,
      ['Recipe', 'Median', 'p95', 'p99', 'p99 / median'],
      [
          ['PLUMERIA NOODLES', '22 min', '28 min', '31 min', '1.41× (stable)'],
          ['LILAC NOODLES', '24 min', '29 min', '33 min', '1.38× (stable)'],
          ['JASMINE NOODLES', '27 min', '41 min', '58 min', '2.15× (variable)'],
          ['GALAXY NOODLES', '31 min', '45 min', '67 min', '2.16× (variable)'],
      ],
      col_widths=[2.0, 0.9, 0.8, 0.8, 1.6])
para(doc,
     'GALAXY is 40 % slower at the median and twice as variable at the tail '
     'compared with PLUMERIA. Production planning currently allocates a flat '
     '30-minute slot per batch — that under-specs GALAXY by 50 %, leading to '
     'cascading line delays.')
lead('What you do with it.')
para(doc,
     'Either re-slot GALAXY at 45 minutes or commission a process improvement '
     'on it. Without this report the line owner has no leverage to ask for '
     'either.')

# --- 8 ----------------------------------------------------------------------
h3(doc, '8. Predictive-vs-actual ROC')
lead('What it shows.')
para(doc,
     'The spike_15m ONNX model emits a probability that a weight-spike will '
     'occur in the next 15 minutes. This report compares those predictions '
     'against what actually happened (from the rule engine\'s ground-truth '
     'labels) and plots a ROC curve. The AUC is the headline KPI for "is our '
     'ML still trustworthy."')
lead('How it is computed.')
code(doc,
     'FOR each prediction P in last 30 days:\n'
     '  truth = (any CRITICAL rule fired within [P.ts, P.ts + 15min])\n'
     '  pairs.append((P.probability, truth))\n'
     'roc = sweep threshold across pairs\n'
     'AUC = area under curve')
lead('Worked example — monthly evolution:')
table(doc,
      ['Month', 'Predictions', 'AUC', 'TPR @ 90% specificity', 'Status'],
      [
          ['March', '218,400', '0.89', '0.71', 'Healthy'],
          ['April', '218,400', '0.87', '0.69', 'Healthy'],
          ['May', '218,400', '0.74', '0.51', 'Drift — retrain ⚠'],
      ],
      col_widths=[1.0, 1.2, 0.7, 1.7, 1.5])
para(doc,
     'May\'s drop is large enough that the model can no longer be trusted at '
     'the 90 %-specificity operating point we previously deployed. Likely '
     'cause: a new recipe (LILAC) was introduced in late April and the '
     'training data does not contain it.')
lead('What you do with it.')
para(doc,
     'Gates whether to enable / disable the inline RISK column from §10.1. '
     'ML governance asks for this curve monthly; without it, ML deployment is '
     'faith-based. Retrain with the last 60 days of data including the new '
     'recipe; re-measure.')

# --- 9 ----------------------------------------------------------------------
h3(doc, '9. Recipe drift detection')
lead('What it shows.')
para(doc,
     'PSM-10 already logs an event whenever an SP changes between two polls. '
     'Aggregating those events per recipe + RM produces an SP-evolution '
     'timeline that exposes silent retuning — an engineer adjusting setpoints '
     'without filing a recipe-change form.')
lead('How it is computed.')
code(doc,
     'FOR each PSM row over period:\n'
     '  sp_history[(recipe, RM)].append((ts, SP))\n'
     'For each (recipe, RM) plot SP vs ts, annotate at every value change.')
lead('Worked example — PLUMERIA / Salt SP across two months:')
table(doc,
      ['Window', 'SP (kg)', 'Source'],
      [
          ['Apr 1 – Apr 10', '10.50', 'baseline'],
          ['Apr 11 → Apr 25', '10.55', 'PSM-10 event logged, no paperwork'],
          ['May 1 → present', '10.65', 'PSM-10 event logged, no paperwork'],
      ],
      col_widths=[2.4, 1.0, 3.0])
para(doc,
     'That is a 1.4 % SP drift in 30 days, no documentation, no recipe '
     'revision. The customer\'s spec calls for ±1 %. The recipe is now out of '
     'spec on paper — but it has been silently moved.')
lead('What you do with it.')
para(doc,
     'This is the audit-defence report. When the QA auditor asks "why does '
     'this batch say SP = 10.65 when the master recipe says 10.50?", the '
     'report tells you exactly when the drift happened so you can roll back '
     'or formalise the change.')

# --- 10 ---------------------------------------------------------------------
h3(doc, '10. Noodle-type purity audit')
lead('What it shows.')
para(doc,
     'For every PSM batch, the fraction of day silos that fed it during its '
     'production window reporting the expected noodle type. Anything below '
     '100 % means the bag potentially contains the wrong noodle type — a '
     'high-impact customer complaint and a regulatory red flag.')
lead('How it is computed.')
code(doc,
     'FOR each PSM batchKey with recipe R and time window [t_start, t_end]:\n'
     '  expected = recipe_to_noodle_map[R]   # e.g. PLUMERIA -> PLUMERIA NOODLES\n'
     '  observed = [silo.noodleType for silo in DAY_SILOS\n'
     '              if any reading in window]\n'
     '  matches  = count(o for o in observed if o == expected)\n'
     '  purity[batchKey] = matches / len(observed) * 100')
lead('Worked example — Batch PSM:01:38:PLUMERIA NOODLES (Apr 26, 14:00–14:22):')
table(doc,
      ['Silo', 'Reported noodle', 'Matches expected?'],
      [
          ['Day_silo_1', 'PLUMERIA NOODLES', '✓'],
          ['Day_silo_2', 'PLUMERIA NOODLES', '✓'],
          ['Day_silo_3', 'JASMINE NOODLES', '✗'],
          ['Day_silo_4', 'PLUMERIA NOODLES', '✓'],
          ['Day_silo_5', 'PLUMERIA NOODLES', '✓'],
          ['Day_silo_6', 'PLUMERIA NOODLES', '✓'],
      ],
      col_widths=[1.4, 2.2, 1.6])
para(doc,
     'Purity = 5 / 6 = 83.3 %. One silo was still carrying the previous '
     'recipe\'s product. The batch is suspect — quarantine for sampling.')
lead('What you do with it.')
para(doc,
     'A "purity %" column joins the per-batch report bundled with shipments. '
     'Regulated customers get a verifiable purity figure; internally, a steady '
     '< 100 % rate on a specific silo becomes a maintenance signal that '
     'reinforces Insight #5.')

# ============================================================================
# 12. KPIs SURFACED
# ============================================================================
page_break(doc)
h1(doc, '12. KPIs Already Surfaced')
para(doc, 'Live numbers on the current dashboard; any of these can be promoted to executive cards:')
for it in [
    'Total rows monitored, total batches monitored',
    'Active filter health score (global, scoped to current view)',
    'Per-batch health score (with passing / total fraction)',
    'Active CRITICAL / WARNING / INFO count by zone',
    'Deviation distribution (green / yellow / pink banding on every numeric row)',
    'Date span of currently loaded data',
    'Per-RM, per-recipe, per-mixer, per-cascade row counts via the dropdowns',
]:
    bullet(doc, it)

# ============================================================================
# 13. TECH STACK
# ============================================================================
h1(doc, '13. Technical Stack')
table(doc,
      ['Layer', 'Technology', 'Notes'],
      [
          ['Frontend', 'Angular 17, TypeScript 5, Tailwind CSS', 'Standalone components, signal-based reactivity'],
          ['Frontend state', 'Angular signals (no NgRx)', 'signal, computed, input.required, output'],
          ['Backend', '.NET 8 Minimal API', 'C# 12, async loaders, in-memory caching'],
          ['Data ingestion', 'ClosedXML (workbook reader)', 'Streams .xlsx files from data/'],
          ['ML runtime', 'Microsoft.ML.OnnxRuntime', 'Loads 6 ONNX models on startup'],
          ['Build / deploy', 'ng build + dotnet publish', 'Localhost; container-ready'],
          ['Repo', 'Git on GitHub — tag-health', 'Branch: master'],
      ],
      col_widths=[1.4, 2.2, 3.3])

# ============================================================================
# 14. RISKS
# ============================================================================
h1(doc, '14. Risks & Mitigations')
table(doc,
      ['Risk', 'Likelihood', 'Mitigation'],
      [
          ['Historian schema changes break parser substring detection', 'Medium', 'Parser is centralised; one file changes. Add a smoke test asserting at least one row per zone parses.'],
          ['ML models drift as recipes change', 'High', 'Add predictive-vs-actual ROC report (Insight #8); retrain monthly.'],
          ['Workbook loading > 5 s on large files', 'Medium', 'Already cached per-zone after first load; long-term move to streaming ingestion.'],
          ['User-action attribution is impossible without operator annotations', 'High', 'Operator annotations item (mid-term #8) closes this gap.'],
      ],
      col_widths=[3.0, 1.0, 3.0])

# ============================================================================
# 15. CLOSING
# ============================================================================
h1(doc, '15. Closing')
para(doc,
     'The platform converts a reactive Excel workflow into a proactive, '
     'rule-validated, ML-augmented operations view. Phase 1 is complete and '
     'demonstrable; the immediate next move is to surface the existing ONNX '
     'predictions inline and to ship the alert-log tab so the system can '
     'replace the daily review meeting rather than supplement it.')
para(doc,
     'Demonstration is available locally — backend on localhost:5050, frontend '
     'on localhost:4200. The repository at github.com/Abhiramdadi17/tag-health '
     'contains everything required to reproduce the build.')

# ============================================================================
# QUICK START
# ============================================================================
h2(doc, 'Appendix — Quick Start')
code(doc,
     '# Backend\n'
     'cd backend\n'
     'dotnet run                # listens on http://localhost:5050\n\n'
     '# Frontend (new terminal)\n'
     'cd frontend\n'
     'npm install\n'
     'npm start                 # serves http://localhost:4200')

# --- save -------------------------------------------------------------------
target = _resolve_out_path(OUT)
doc.save(target)
print(f'wrote {target}')
if target != OUT:
    print(f'(original {OUT.name} is locked - close Word and rename {target.name} to {OUT.name})')
